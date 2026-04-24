import os

try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml, OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

TABLE_HEADER_BG = 'E8EEF5'


def _rfonts(style_obj, cn, en):
    rPr = style_obj.element.get_or_add_rPr()
    rf = rPr.find(qn('w:rFonts'))
    if rf is None:
        rf = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rf)
    rf.set(qn('w:ascii'), en)
    rf.set(qn('w:hAnsi'), en)
    rf.set(qn('w:eastAsia'), cn)


def _spacing(pPr, line=None, before=None, after=None):
    sp = pPr.find(qn('w:spacing'))
    if sp is None:
        sp = parse_xml(f'<w:spacing {nsdecls("w")}/>')
        pPr.append(sp)
    if line is not None:
        sp.set(qn('w:line'), str(int(line * 20)))
        sp.set(qn('w:lineRule'), 'exact')
    if before is not None:
        sp.set(qn('w:before'), str(int(before * 20)))
    if after is not None:
        sp.set(qn('w:after'), str(int(after * 20)))


def _no_indent(pPr):
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = parse_xml(f'<w:ind {nsdecls("w")}/>')
        pPr.append(ind)
    ind.set(qn('w:firstLine'), '0')


def _setup_page(section):
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3.73)
    section.bottom_margin = Cm(3.53)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)
    section.header_distance = Cm(2.0)
    section.footer_distance = Cm(1.7)


def _set_run_font(run, cn, en, sz, bold=False):
    run.font.size = Pt(sz)
    run.font.name = en
    run.font.bold = bold
    rPr = run._element.get_or_add_rPr()
    rf = rPr.find(qn('w:rFonts'))
    if rf is None:
        rf = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rf)
    rf.set(qn('w:eastAsia'), cn)


class DocxBuilder:

    def __init__(self, report_title, project_name, report_date,
                 author='XXX', classification='内部', header_text=None):
        self._title = report_title
        self._project = project_name
        self._date = report_date
        self._author = author
        self._classification = classification
        self._header_text = header_text or report_title
        self._chapter = 0
        self._table_in_chapter = 0
        self._doc = None

    def _next_table_num(self):
        self._table_in_chapter += 1
        return self._table_in_chapter

    def _new_chapter(self):
        self._chapter += 1
        self._table_in_chapter = 0

    def _setup_styles(self):
        doc = self._doc
        normal = doc.styles['Normal']
        normal.font.name = 'Times New Roman'
        normal.font.size = Pt(10.5)
        _rfonts(normal, 'SimSun', 'Times New Roman')
        pPr = normal.element.get_or_add_pPr()
        _spacing(pPr, line=28)
        ind = pPr.find(qn('w:ind'))
        if ind is None:
            ind = parse_xml(f'<w:ind {nsdecls("w")}/>')
            pPr.append(ind)
        ind.set(qn('w:firstLine'), '420')

        for name, sz, bef, aft in [
            ('Heading 1', 16, 24, 12),
            ('Heading 2', 14, 12, 6),
            ('Heading 3', 12, 6, 6),
        ]:
            s = doc.styles[name]
            _rfonts(s, 'SimHei', 'Times New Roman')
            s.font.size = Pt(sz)
            s.font.bold = True
            pPr = s.element.get_or_add_pPr()
            pPr.append(parse_xml(f'<w:keepNext {nsdecls("w")}/>'))
            pPr.append(parse_xml(f'<w:keepLines {nsdecls("w")}/>'))
            _no_indent(pPr)
            _spacing(pPr, line=28, before=bef, after=aft)

        cap = doc.styles.add_style('ReportCaption', WD_STYLE_TYPE.PARAGRAPH)
        cap.base_style = normal
        _rfonts(cap, 'SimSun', 'Times New Roman')
        cap.font.size = Pt(9)
        pPr = cap.element.get_or_add_pPr()
        pPr.append(parse_xml(f'<w:jc {nsdecls("w")} w:val="center"/>'))
        pPr.append(parse_xml(f'<w:keepLines {nsdecls("w")}/>'))
        _no_indent(pPr)
        _spacing(pPr, line=18, before=12, after=12)

        tc = doc.styles.add_style('TableCellPara', WD_STYLE_TYPE.PARAGRAPH)
        _rfonts(tc, 'SimSun', 'Times New Roman')
        tc.font.size = Pt(9)
        pPr = tc.element.get_or_add_pPr()
        _no_indent(pPr)
        _spacing(pPr, line=20, before=2, after=2)

        thc = doc.styles.add_style('TableHeaderPara', WD_STYLE_TYPE.PARAGRAPH)
        thc.base_style = tc
        _rfonts(thc, 'SimHei', 'Times New Roman')
        thc.font.size = Pt(10)
        thc.font.bold = True
        pPr = thc.element.get_or_add_pPr()
        pPr.append(parse_xml(f'<w:jc {nsdecls("w")} w:val="center"/>'))

    def _add_cover(self):
        doc = self._doc
        _setup_page(doc.sections[0])

        for _ in range(6):
            doc.add_paragraph()

        p = doc.add_paragraph(self._title)
        p.paragraph_format.first_line_indent = Cm(0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(20)
        for r in p.runs:
            _set_run_font(r, 'SimHei', 'Times New Roman', 22, bold=True)

        p = doc.add_paragraph(f'{self._project} 项目')
        p.paragraph_format.first_line_indent = Cm(0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(60)
        for r in p.runs:
            _set_run_font(r, 'SimSun', 'Times New Roman', 12)

        for t in [f'分析人员：{self._author}', f'报告日期：{self._date}', f'密级：{self._classification}']:
            p = doc.add_paragraph(t)
            p.paragraph_format.first_line_indent = Cm(0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(6)

        doc.add_section()
        s2 = doc.sections[1]
        _setup_page(s2)

        hp = s2.header.paragraphs[0]
        hp.text = ''
        hr = hp.add_run(self._header_text)
        _set_run_font(hr, 'SimSun', 'Times New Roman', 9)
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        fp = s2.footer.paragraphs[0]
        fp.text = ''
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for tag, val in [('begin', None), (None, ' PAGE '), ('end', None)]:
            r = fp.add_run()
            if tag:
                fc = OxmlElement('w:fldChar')
                fc.set(qn('w:fldCharType'), tag)
                r._element.append(fc)
            else:
                it = OxmlElement('w:instrText')
                it.set(qn('xml:space'), 'preserve')
                it.text = val
                r._element.append(it)

    def new_chapter(self):
        self._new_chapter()

    def add_heading(self, text, level):
        return self._doc.add_paragraph(text, style=f'Heading {level}')

    def add_body(self, text):
        return self._doc.add_paragraph(text, style='Normal')

    def add_caption(self, text):
        return self._doc.add_paragraph(text, style='ReportCaption')

    def table_caption(self, desc):
        num = self._next_table_num()
        return self.add_caption(f'表 {self._chapter}-{num}  {desc}')

    def add_table(self, headers, rows):
        doc = self._doc
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER

        tbl = table._tbl
        tblPr = tbl.tblPr
        b = tblPr.find(qn('w:tblBorders'))
        if b is not None:
            tblPr.remove(b)
        tblPr.append(parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'<w:top w:val="single" w:sz="12" w:color="000000" w:space="0"/>'
            f'<w:bottom w:val="single" w:sz="12" w:color="000000" w:space="0"/>'
            f'<w:left w:val="nil"/><w:right w:val="nil"/>'
            f'<w:insideH w:val="nil"/><w:insideV w:val="nil"/>'
            f'</w:tblBorders>'))
        m = tblPr.find(qn('w:tblCellMar'))
        if m is not None:
            tblPr.remove(m)
        tblPr.append(parse_xml(
            f'<w:tblCellMar {nsdecls("w")}>'
            f'<w:top w:w="60" w:type="dxa"/><w:left w:w="100" w:type="dxa"/>'
            f'<w:bottom w:w="60" w:type="dxa"/><w:right w:w="100" w:type="dxa"/>'
            f'</w:tblCellMar>'))

        hdr = table.rows[0]
        hdr._tr.get_or_add_trPr().append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))
        for i, h in enumerate(headers):
            cell = hdr.cells[i]
            cell.text = ''
            p = cell.paragraphs[0]
            p.style = doc.styles['TableHeaderPara']
            p.add_run(h)
            tcPr = cell._tc.get_or_add_tcPr()
            tcPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{TABLE_HEADER_BG}" w:val="clear"/>'))
            tcPr.append(parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                f'<w:bottom w:val="single" w:sz="6" w:color="000000" w:space="0"/>'
                f'</w:tcBorders>'))

        for r_idx, row_data in enumerate(rows):
            for c_idx, val in enumerate(row_data):
                cell = table.rows[1 + r_idx].cells[c_idx]
                cell.text = ''
                p = cell.paragraphs[0]
                p.style = doc.styles['TableCellPara']
                if c_idx == 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run(str(val))
        return table

    def build(self):
        self._doc = Document()
        self._setup_styles()
        self._add_cover()
        return self._doc

    def save(self, output_path):
        output_path = os.path.abspath(output_path)
        if self._doc is None:
            self.build()
        self._doc.save(output_path)
        print(f'报告已生成：{output_path}')
        return output_path
